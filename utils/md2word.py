import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement as docx_OxmlElement
import re
import os
from bs4 import BeautifulSoup
from datetime import datetime

def markdown_to_word(markdown_content, output_filename, company_name="Medtronic", author_name=""):
    """
    Convert markdown content to a formatted Word document
    """
    # Convert markdown to HTML
    html = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')
    
    # Create a new Word document
    doc = Document()
    
    # Set document properties
    doc.core_properties.author = author_name
    doc.core_properties.title = "SPM Solution Design Document"
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header with company name
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(company_name)
    header_run.bold = True
    header_run.font.size = Pt(10)
    header_run.font.name = 'Calibri'
    
    # Add footer with page numbers and date
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add document creation date to left side of footer
    current_date = datetime.now().strftime("%B %d, %Y")
    date_run = footer_para.add_run(current_date)
    date_run.font.size = Pt(9)
    date_run.font.name = 'Calibri'
    
    # Add tab and page number
    footer_para.add_run("      Page ")
    page_num = footer_para.add_run()
    page_num.font.size = Pt(9)
    page_num.font.name = 'Calibri'
    
    # Add field for page numbers
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = "PAGE"
    
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    
    page_num._element.append(fld_char)
    page_num._element.append(instr_text)
    page_num._element.append(fld_char2)
    
    # Add total pages
    footer_para.add_run(" of ")
    total_pages = footer_para.add_run()
    total_pages.font.size = Pt(9)
    total_pages.font.name = 'Calibri'
    
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = "NUMPAGES"
    
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    
    total_pages._element.append(fld_char)
    total_pages._element.append(instr_text)
    total_pages._element.append(fld_char2)
    
    # Define heading styles
    styles = {
        'h1': {'font_size': Pt(18), 'bold': True, 'color': RGBColor(0, 51, 153), 'space_before': Pt(24), 'space_after': Pt(6)},
        'h2': {'font_size': Pt(16), 'bold': True, 'color': RGBColor(0, 51, 153), 'space_before': Pt(18), 'space_after': Pt(6)},
        'h3': {'font_size': Pt(14), 'bold': True, 'color': RGBColor(0, 51, 153), 'space_before': Pt(12), 'space_after': Pt(6)},
        'h4': {'font_size': Pt(12), 'bold': True, 'color': RGBColor(0, 51, 153), 'space_before': Pt(10), 'space_after': Pt(4)},
    }
    
    # Extract title (the first h1) for cover page
    title = soup.find('h1').text if soup.find('h1') else "SPM Solution Design Document"
    subtitle = soup.find('h2').text if soup.find('h2') else ""
    
    # Create cover page
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_para.paragraph_format.space_after = Pt(120)
    
    # Add logo placeholder (would need to be replaced with actual logo)
    logo_run = cover_para.add_run("[Company Logo]\n\n")
    logo_run.font.size = Pt(11)
    logo_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add title
    title_run = cover_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 51, 153)
    title_run.font.name = 'Calibri'
    
    # Add subtitle if available
    if subtitle:
        cover_para.add_run("\n\n")
        subtitle_run = cover_para.add_run(subtitle)
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(20)
        subtitle_run.font.color.rgb = RGBColor(0, 51, 153)
        subtitle_run.font.name = 'Calibri'
    
    # Add date and author
    cover_para.add_run("\n\n\n")
    date_author_run = cover_para.add_run(f"{current_date}\n\n{author_name if author_name else '[Author Name]'}")
    date_author_run.font.size = Pt(12)
    date_author_run.font.name = 'Calibri'
    
    # Add confidentiality statement
    cover_para.add_run("\n\n\n")
    confidential_run = cover_para.add_run("CONFIDENTIAL")
    confidential_run.bold = True
    confidential_run.font.size = Pt(12)
    confidential_run.font.color.rgb = RGBColor(192, 0, 0)
    confidential_run.font.name = 'Calibri'
    
    # Add page break after cover page
    doc.add_page_break()
    
    # Table of contents
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run("Table of Contents")
    toc_run.bold = True
    toc_run.font.size = Pt(16)
    toc_run.font.color.rgb = RGBColor(0, 51, 153)
    toc_run.font.name = 'Calibri'
    
    # Add TOC field
    para = doc.add_paragraph()
    run = para.add_run()
    
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fld_char)
    run._element.append(instr_text)
    run._element.append(fld_char2)
    run._element.append(fld_char3)
    
    # Add page break after TOC
    doc.add_page_break()
    
    # Process HTML elements
    elements = list(soup.children)
    
    # Track if we're in a list
    in_list = False
    list_type = None
    list_items = []
    
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol', 'li', 'pre', 'hr']):
        # Handle headings
        if element.name in ['h1', 'h2', 'h3', 'h4']:
            # Check if we need to end a list
            if in_list:
                add_list_to_doc(doc, list_items, list_type)
                in_list = False
                list_items = []
            
            # Get heading level (1-4)
            level = int(element.name[1])
            
            # Add the heading with the appropriate style
            para = doc.add_paragraph()
            para.style = f'Heading {level}'
            run = para.add_run(element.text)
            run.bold = styles[element.name]['bold']
            run.font.size = styles[element.name]['font_size']
            run.font.color.rgb = styles[element.name]['color']
            run.font.name = 'Calibri'
            
            # Add appropriate spacing
            para.paragraph_format.space_before = styles[element.name]['space_before']
            para.paragraph_format.space_after = styles[element.name]['space_after']
            
        # Handle paragraphs
        elif element.name == 'p':
            # Check if we need to end a list
            if in_list:
                add_list_to_doc(doc, list_items, list_type)
                in_list = False
                list_items = []
            
            para = doc.add_paragraph()
            
            # Process inline elements
            for child in element.children:
                if hasattr(child, 'name') and child.name == 'strong':
                    run = para.add_run(child.text)
                    run.bold = True
                elif hasattr(child, 'name') and child.name == 'em':
                    run = para.add_run(child.text)
                    run.italic = True
                elif hasattr(child, 'name') and child.name == 'code':
                    run = para.add_run(child.text)
                    run.font.name = 'Courier New'
                else:
                    run = para.add_run(str(child))
                
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
            
        # Handle tables
        elif element.name == 'table':
            # Check if we need to end a list
            if in_list:
                add_list_to_doc(doc, list_items, list_type)
                in_list = False
                list_items = []
            
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table_cell = table.cell(i, j)
                        table_cell.text = cell.get_text().strip()
                        
                        # Make header row bold
                        if i == 0 or cell.name == 'th':
                            for paragraph in table_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                
                # Add space after the table
                doc.add_paragraph()
        
        # Handle lists
        elif element.name in ['ul', 'ol']:
            in_list = True
            list_type = 'unordered' if element.name == 'ul' else 'ordered'
        
        elif element.name == 'li' and in_list:
            list_items.append(element.text)
        
        # Handle code blocks
        elif element.name == 'pre':
            # Check if we need to end a list
            if in_list:
                add_list_to_doc(doc, list_items, list_type)
                in_list = False
                list_items = []
            
            code_para = doc.add_paragraph()
            code_text = element.get_text()
            code_run = code_para.add_run(code_text)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
            
            # Add light gray shading to code blocks
            code_para.paragraph_format.left_indent = Inches(0.5)
            code_para.paragraph_format.right_indent = Inches(0.5)
            
        # Handle horizontal rule
        elif element.name == 'hr':
            # Check if we need to end a list
            if in_list:
                add_list_to_doc(doc, list_items, list_type)
                in_list = False
                list_items = []
            
            doc.add_paragraph('_' * 80)
    
    # Add any remaining list
    if in_list:
        add_list_to_doc(doc, list_items, list_type)
    
    # Save the document
    doc.save(output_filename)
    print(f"Document saved as {output_filename}")

def add_list_to_doc(doc, items, list_type):
    """Add a list to the document"""
    for i, item in enumerate(items):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(0.0)
        
        # Add bullet or number
        if list_type == 'unordered':
            para.style = 'List Bullet'
        else:
            para.style = 'List Number'
        
        run = para.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

def OxmlElement(tag):
    """Create an OOXML element with the given tag"""
    from docx.oxml.ns import nsdecls
    return docx_OxmlElement(tag)

if __name__ == "__main__":
    # Set up paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    input_file = os.path.join(script_dir, "medtronic_spm_design.md")
    output_file = os.path.join(script_dir, "Medtronic_SPM_Solution_Design.docx")
    
    # Check if input file exists, if not, create it with a prompt
    if not os.path.exists(input_file):
        print(f"Input file {input_file} not found. Please create this file with your Markdown content.")
        # You could also add code here to get the markdown content from the clipboard or another source
    else:
        # Read markdown content
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert to Word
        markdown_to_word(markdown_content, output_file, company_name="Medtronic", author_name="Your Name")
